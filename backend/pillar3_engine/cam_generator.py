import logging
import asyncio
import os
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BOLD
from docx.enum.table import WD_TABLE_ALIGNMENT
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

from shared.ollama_client import ollama_client
from shared.models import (
    CAMDocument, CreditDecision, FiveCScores, CreditGrade,
    ExtractedData, ValidationFlag, Severity, CIBILData,
    LoanLimitResult, ResearchReport
)
from config import config

logger = logging.getLogger(__name__)

class CAMGenerator:
    """Generate Credit Appraisal Memo (CAM) in Word and PDF formats"""
    
    def __init__(self):
        self.output_dir = Path(config.CAM_OUTPUT_DIR)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # System prompts for different sections
        self.section_prompts = {
            "executive_summary": """You are a senior credit analyst writing an executive summary for a credit committee.
Write a concise 3-line summary including: 1) Credit decision, 2) Loan limit in ₹ Crores, 3) Key reason for decision.
Use formal Indian banking English. Be specific with figures. Maximum 50 words.""",
            
            "character": """You are a senior credit analyst writing the Character section of a Credit Appraisal Memo.
Analyze the borrower's character based on compliance history, credit bureau data, and management quality.
Use formal Indian banking English. Cite specific data points. Maximum 200 words.""",
            
            "capacity": """You are a senior credit analyst writing the Capacity section of a Credit Appraisal Memo.
Analyze the borrower's capacity to repay based on cash flows, DSCR, and operational efficiency.
Use formal Indian banking English. Include specific financial ratios. Maximum 200 words.""",
            
            "capital": """You are a senior credit analyst writing the Capital section of a Credit Appraisal Memo.
Analyze the borrower's capital structure, net worth, and financial stability.
Use formal Indian banking English. Reference specific balance sheet metrics. Maximum 200 words.""",
            
            "collateral": """You are a senior credit analyst writing the Collateral section of a Credit Appraisal Memo.
Analyze the collateral offered, its valuation, and security coverage.
Use formal Indian banking English. Include collateral-to-loan ratios. Maximum 200 words.""",
            
            "conditions": """You are a senior credit analyst writing the Conditions section of a Credit Appraisal Memo.
Analyze external factors including sector conditions, regulatory environment, and market risks.
Use formal Indian banking English. Reference specific sector data. Maximum 200 words.""",
            
            "risk_flags": """You are a senior credit analyst summarizing risk flags for a credit committee.
Summarize all validation flags and risk findings in order of severity.
Use formal Indian banking English. Be specific about implications. Maximum 150 words."""
        }
    
    async def generate_cam(
        self,
        credit_decision: CreditDecision,
        extracted_data: List[ExtractedData],
        validation_flags: List[ValidationFlag],
        research_report: Optional[ResearchReport] = None,
        officer_observations: Optional[str] = None,
        shap_chart_path: Optional[str] = None
    ) -> CAMDocument:
        """
        Generate complete CAM document in Word and PDF formats
        
        Args:
            credit_decision: Final credit decision with scores
            extracted_data: All extracted document data
            validation_flags: Cross-validation flags
            research_report: External research findings
            officer_observations: Field visit observations
            shap_chart_path: Path to SHAP explanation chart
            
        Returns:
            CAMDocument with file paths and metadata
        """
        try:
            logger.info(f"Starting CAM generation for {credit_decision.company_name}")
            
            # Generate timestamp and file paths
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            base_filename = f"cam_{credit_decision.cin}_{timestamp}"
            
            word_path = self.output_dir / f"{base_filename}.docx"
            pdf_path = self.output_dir / f"{base_filename}.pdf"
            
            # Generate narratives for each section
            narratives = await self._generate_narratives(
                credit_decision, extracted_data, validation_flags, research_report
            )
            
            # Generate Word document
            await self._generate_word_document(
                word_path, credit_decision, narratives, extracted_data,
                validation_flags, research_report, officer_observations, shap_chart_path
            )
            
            # Generate PDF document
            await self._generate_pdf_document(
                pdf_path, credit_decision, narratives, extracted_data,
                validation_flags, research_report, officer_observations, shap_chart_path
            )
            
            cam_document = CAMDocument(
                word_path=str(word_path),
                pdf_path=str(pdf_path),
                decision=credit_decision,
                research_report=research_report,
                extracted_data=extracted_data,
                officer_observations=officer_observations
            )
            
            logger.info(f"CAM generated successfully: {word_path.name}")
            return cam_document
            
        except Exception as e:
            logger.error(f"CAM generation failed: {e}")
            raise
    
    async def _generate_narratives(
        self,
        credit_decision: CreditDecision,
        extracted_data: List[ExtractedData],
        validation_flags: List[ValidationFlag],
        research_report: Optional[ResearchReport] = None
    ) -> Dict[str, str]:
        """Generate AI-powered narratives for each CAM section"""
        try:
            narratives = {}
            
            # Prepare context data for Ollama
            context_data = self._prepare_context_data(
                credit_decision, extracted_data, validation_flags, research_report
            )
            
            # Generate each section narrative
            for section, system_prompt in self.section_prompts.items():
                try:
                    prompt = self._build_section_prompt(section, context_data)
                    narrative = await ollama_client.generate(
                        prompt=prompt,
                        system=system_prompt,
                        temperature=0.1
                    )
                    narratives[section] = narrative.strip()
                    
                except Exception as e:
                    logger.warning(f"Failed to generate {section} narrative: {e}")
                    narratives[section] = f"[{section.replace('_', ' ').title()} section pending manual completion]"
            
            return narratives
            
        except Exception as e:
            logger.error(f"Narrative generation failed: {e}")
            # Return placeholder narratives
            return {section: f"[{section.replace('_', ' ').title()} section pending manual completion]" 
                   for section in self.section_prompts.keys()}
    
    def _prepare_context_data(
        self,
        credit_decision: CreditDecision,
        extracted_data: List[ExtractedData],
        validation_flags: List[ValidationFlag],
        research_report: Optional[ResearchReport] = None
    ) -> Dict[str, Any]:
        """Prepare context data for narrative generation"""
        context = {
            "company_name": credit_decision.company_name,
            "cin": credit_decision.cin,
            "total_score": credit_decision.total_score,
            "grade": credit_decision.grade.value,
            "five_c_scores": {
                "character": credit_decision.five_c_scores.character,
                "capacity": credit_decision.five_c_scores.capacity,
                "capital": credit_decision.five_c_scores.capital,
                "collateral": credit_decision.five_c_scores.collateral,
                "conditions": credit_decision.five_c_scores.conditions
            },
            "validation_flags": [
                {
                    "type": flag.flag_type,
                    "severity": flag.severity.value,
                    "description": flag.description
                }
                for flag in validation_flags
            ]
        }
        
        # Add loan structure if available
        if credit_decision.loan_limit:
            context["loan_limit_cr"] = credit_decision.loan_limit.final_limit_inr / 10000000
            context["interest_rate"] = credit_decision.interest_rate_pct
            context["tenor_months"] = credit_decision.tenor_months
        
        # Add CIBIL data if available
        if credit_decision.cibil_data:
            context["cibil"] = {
                "cmr_rank": credit_decision.cibil_data.cmr_rank,
                "overdue_amount": credit_decision.cibil_data.overdue_amount_inr,
                "enquiries_6m": credit_decision.cibil_data.credit_enquiries_6m
            }
        
        # Add financial data
        for data in extracted_data:
            if data.financial_fields:
                context["financial"] = {
                    "revenue_inr": data.financial_fields.revenue_inr,
                    "ebitda_inr": data.financial_fields.ebitda_inr,
                    "pat_inr": data.financial_fields.pat_inr,
                    "net_worth_inr": data.financial_fields.net_worth_inr,
                    "total_debt_inr": data.financial_fields.total_debt_inr
                }
                break
        
        return context
    
    def _build_section_prompt(self, section: str, context_data: Dict[str, Any]) -> str:
        """Build section-specific prompt with context data"""
        
        if section == "executive_summary":
            return f"""Based on the following credit analysis, write an executive summary:

Company: {context_data['company_name']} (CIN: {context_data['cin']})
Total Score: {context_data['total_score']:.1f}/100
Grade: {context_data['grade']}
Loan Limit: ₹{context_data.get('loan_limit_cr', 'N/A')} Cr
Interest Rate: {context_data.get('interest_rate', 'N/A')}%

Five Cs Scores:
- Character: {context_data['five_c_scores']['character']:.1f}
- Capacity: {context_data['five_c_scores']['capacity']:.1f}
- Capital: {context_data['five_c_scores']['capital']:.1f}
- Collateral: {context_data['five_c_scores']['collateral']:.1f}
- Conditions: {context_data['five_c_scores']['conditions']:.1f}

Risk Flags: {len(context_data['validation_flags'])} identified"""

        elif section == "character":
            prompt = f"""Character Analysis for {context_data['company_name']}:

Character Score: {context_data['five_c_scores']['character']:.1f}/100"""
            
            if 'cibil' in context_data:
                prompt += f"""
CIBIL Data:
- CMR Rank: {context_data['cibil']['cmr_rank']}/10
- Overdue Amount: ₹{context_data['cibil']['overdue_amount']:,.0f}
- Credit Enquiries (6m): {context_data['cibil']['enquiries_6m']}"""
            
            return prompt
        
        elif section == "capacity":
            prompt = f"""Capacity Analysis for {context_data['company_name']}:

Capacity Score: {context_data['five_c_scores']['capacity']:.1f}/100"""
            
            if 'financial' in context_data:
                prompt += f"""
Financial Metrics:
- EBITDA: ₹{context_data['financial']['ebitda_inr']:,.0f}
- Total Debt: ₹{context_data['financial']['total_debt_inr']:,.0f}"""
            
            return prompt
        
        elif section == "risk_flags":
            flag_text = "\n".join([
                f"- {flag['type']} ({flag['severity']}): {flag['description']}"
                for flag in context_data['validation_flags']
            ])
            return f"""Risk Flags Summary for {context_data['company_name']}:

{flag_text}"""
        
        # Generic prompt for other sections
        return f"""{section.replace('_', ' ').title()} Analysis for {context_data['company_name']}:

Score: {context_data['five_c_scores'].get(section, 'N/A')}/100
Total Credit Score: {context_data['total_score']:.1f}/100
Grade: {context_data['grade']}"""
    
    async def _generate_word_document(
        self,
        output_path: Path,
        credit_decision: CreditDecision,
        narratives: Dict[str, str],
        extracted_data: List[ExtractedData],
        validation_flags: List[ValidationFlag],
        research_report: Optional[ResearchReport] = None,
        officer_observations: Optional[str] = None,
        shap_chart_path: Optional[str] = None
    ):
        """Generate Word document"""
        try:
            doc = Document()
            
            # Add title
            title = doc.add_paragraph('CREDIT APPRAISAL MEMORANDUM')
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.runs[0].bold = True
            title.runs[0].font.size = Pt(16)
            
            # Add header info
            doc.add_paragraph(f'Vivriti Capital')
            doc.add_paragraph(f'Date: {datetime.now().strftime("%d %B %Y")}')
            doc.add_paragraph(f'Borrower: {credit_decision.company_name}')
            doc.add_paragraph(f'CIN: {credit_decision.cin}')
            doc.add_paragraph('')  # Spacer
            
            # Executive Summary
            doc.add_heading('Executive Summary', level=1)
            doc.add_paragraph(narratives.get('executive_summary', '[Executive summary pending]'))
            doc.add_paragraph('')
            
            # Five Cs sections
            sections = [
                ('Character', 'character'),
                ('Capacity', 'capacity'),
                ('Capital', 'capital'),
                ('Collateral', 'collateral'),
                ('Conditions', 'conditions')
            ]
            
            for section_name, section_key in sections:
                doc.add_heading(section_name, level=1)
                doc.add_paragraph(narratives.get(section_key, f'[{section_name} analysis pending]'))
                
                # Add score badge
                score = getattr(credit_decision.five_c_scores, section_key)
                score_para = doc.add_paragraph(f'Score: {score:.1f}/100')
                score_para.runs[0].bold = True
                
                # Color code based on score
                if score >= 80:
                    score_para.runs[0].font.color.rgb = RGBColor(0, 128, 0)  # Green
                elif score >= 60:
                    score_para.runs[0].font.color.rgb = RGBColor(255, 165, 0)  # Orange
                else:
                    score_para.runs[0].font.color.rgb = RGBColor(255, 0, 0)  # Red
                
                doc.add_paragraph('')
            
            # Risk Flags
            if validation_flags:
                doc.add_heading('Risk Flags', level=1)
                doc.add_paragraph(narratives.get('risk_flags', '[Risk analysis pending]'))
                
                # Add risk flags table
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Flag Type'
                hdr_cells[1].text = 'Severity'
                hdr_cells[2].text = 'Description'
                
                for flag in validation_flags:
                    row_cells = table.add_row().cells
                    row_cells[0].text = flag.flag_type
                    row_cells[1].text = flag.severity.value
                    row_cells[2].text = flag.description
                
                doc.add_paragraph('')
            
            # Loan Structure
            if credit_decision.loan_limit:
                doc.add_heading('Loan Structure', level=1)
                loan_table = doc.add_table(rows=4, cols=2)
                loan_table.style = 'Table Grid'
                
                loan_data = [
                    ('Loan Limit', f"₹{credit_decision.loan_limit.final_limit_inr/10000000:.2f} Cr"),
                    ('Interest Rate', f"{credit_decision.interest_rate_pct:.1f}%"),
                    ('Tenor', f"{credit_decision.tenor_months} months"),
                    ('Binding Ceiling', credit_decision.loan_limit.binding_ceiling.replace('_', ' ').title())
                ]
                
                for i, (label, value) in enumerate(loan_data):
                    loan_table.cell(i, 0).text = label
                    loan_table.cell(i, 1).text = value
            
            # Officer Observations
            if officer_observations:
                doc.add_heading('Field Observations', level=1)
                obs_para = doc.add_paragraph(officer_observations)
                obs_para.italic = True
            
            # SHAP Chart
            if shap_chart_path and os.path.exists(shap_chart_path):
                doc.add_heading('Model Explainability', level=1)
                doc.add_paragraph('SHAP analysis provides insights into how different factors influenced the credit decision:')
                
                # Add SHAP summary plot
                doc.add_picture(shap_chart_path, width=Inches(6))
                
                # Add SHAP explanations
                if 'explanations' in shap_data:
                    explanations = shap_data['explanations']
                    for explanation in explanations:
                        p = doc.add_paragraph()
                        p.add_run(f"{explanation['component']} Score: ", bold=True)
                        p.add_run(f"{explanation['score']:.1f}/100")
                        
                        impact_color = '00FF00' if explanation['impact'] == 'positive' else 'FF0000'
                        p.add_run(f" ({explanation['impact']} impact)", color=impact_color)
                        
                        p = doc.add_paragraph()
                        p.add_run("SHAP Impact: ", bold=True)
                        p.add_run(f"{explanation['shap_value']:+.2f}")
                        p.add_run(f" - {explanation['explanation']}")
                
                # Add feature importance
                if 'charts' in shap_data and 'feature_importance_plot' in shap_data['charts']:
                    doc.add_heading('Feature Importance', level=2)
                    doc.add_picture(shap_data['charts']['feature_importance_plot'], width=Inches(5))
                    
                    # Add feature importance table
                    feature_importance = [
                        ['Financial Metrics', 'Revenue, EBITDA, PAT, Net Worth, Total Debt'],
                        ['Credit Indicators', 'CIBIL Rank, Overdue Amount, Credit Enquiries'],
                        ['Operational Metrics', 'Fixed Assets, Cash Balance, GST Turnover'],
                        ['Risk Factors', 'Validation Flags, Sector Risk Score']
                    ]
                    
                    table = doc.add_table(rows=len(feature_importance)+1, cols=2)
                    table.style = 'Table Grid'
                    
                    # Header row
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Category'
                    hdr_cells[1].text = 'Features'
                    for cell in hdr_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                    
                    # Data rows
                    for i, (category, features) in enumerate(feature_importance):
                        row_cells = table.rows[i+1].cells
                        row_cells[0].text = category
                        row_cells[1].text = features
                        for cell in row_cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(9)
                    
                    doc.add_table(table)
            
            # Save document
            doc.save(output_path)
            logger.info(f"Word document saved: {output_path}")
            
        except Exception as e:
            logger.error(f"Word document generation failed: {e}")
            raise
    
    async def _generate_pdf_document(
        self,
        output_path: Path,
        credit_decision: CreditDecision,
        narratives: Dict[str, str],
        extracted_data: List[ExtractedData],
        validation_flags: List[ValidationFlag],
        research_report: Optional[ResearchReport] = None,
        officer_observations: Optional[str] = None,
        shap_chart_path: Optional[str] = None
    ):
        """Generate PDF document using ReportLab"""
        try:
            doc = SimpleDocTemplate(str(output_path), pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                spaceAfter=30,
                alignment=TA_CENTER
            )
            
            # Title
            story.append(Paragraph("CREDIT APPRAISAL MEMORANDUM", title_style))
            story.append(Spacer(1, 12))
            
            # Header info
            header_data = [
                ['Vivriti Capital', ''],
                [f'Date: {datetime.now().strftime("%d %B %Y")}', ''],
                [f'Borrower: {credit_decision.company_name}', ''],
                [f'CIN: {credit_decision.cin}', ''],
            ]
            
            header_table = Table(header_data, colWidths=[4*inch, 2*inch])
            header_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ]))
            story.append(header_table)
            story.append(Spacer(1, 20))
            
            # Executive Summary
            story.append(Paragraph("Executive Summary", styles['Heading2']))
            story.append(Paragraph(narratives.get('executive_summary', '[Executive summary pending]'), styles['Normal']))
            story.append(Spacer(1, 12))
            
            # Five Cs sections
            sections = [
                ('Character', 'character'),
                ('Capacity', 'capacity'),
                ('Capital', 'capital'),
                ('Collateral', 'collateral'),
                ('Conditions', 'conditions')
            ]
            
            for section_name, section_key in sections:
                story.append(Paragraph(section_name, styles['Heading2']))
                story.append(Paragraph(narratives.get(section_key, f'[{section_name} analysis pending]'), styles['Normal']))
                
                # Add score
                score = getattr(credit_decision.five_c_scores, section_key)
                score_color = colors.green if score >= 80 else colors.orange if score >= 60 else colors.red
                story.append(Paragraph(f"<font color='{score_color.hexval()[1:]}'><b>Score: {score:.1f}/100</b></font>", styles['Normal']))
                story.append(Spacer(1, 12))
            
            # Risk Flags
            if validation_flags:
                story.append(Paragraph("Risk Flags", styles['Heading2']))
                story.append(Paragraph(narratives.get('risk_flags', '[Risk analysis pending]'), styles['Normal']))
                
                # Risk flags table
                flag_data = [['Flag Type', 'Severity', 'Description']]
                for flag in validation_flags:
                    flag_data.append([flag.flag_type, flag.severity.value, flag.description])
                
                flag_table = Table(flag_data, colWidths=[2*inch, 1*inch, 3*inch])
                flag_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(flag_table)
                story.append(Spacer(1, 12))
            
            # Loan Structure
            if credit_decision.loan_limit:
                story.append(Paragraph("Loan Structure", styles['Heading2']))
                loan_data = [
                    ['Loan Limit', f"₹{credit_decision.loan_limit.final_limit_inr/10000000:.2f} Cr"],
                    ['Interest Rate', f"{credit_decision.interest_rate_pct:.1f}%"],
                    ['Tenor', f"{credit_decision.tenor_months} months"],
                    ['Binding Ceiling', credit_decision.loan_limit.binding_ceiling.replace('_', ' ').title()]
                ]
                
                loan_table = Table(loan_data, colWidths=[2*inch, 2*inch])
                loan_table.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(loan_table)
                story.append(Spacer(1, 12))
            
            # Officer Observations
            if officer_observations:
                story.append(Paragraph("Field Observations", styles['Heading2']))
                story.append(Paragraph(f"<i>{officer_observations}</i>", styles['Normal']))
            
            # Build PDF
            doc.build(story)
            logger.info(f"PDF document saved: {output_path}")
            
        except Exception as e:
            logger.error(f"PDF document generation failed: {e}")
            raise

# Global instance
cam_generator = CAMGenerator()
